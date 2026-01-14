import os
import re
import fitz  # PyMuPDF
import pdfplumber
import docx
import pandas as pd
import openpyxl
import base64
from typing import List, Dict, Any, Tuple
from tools.doc_type import DocType
from backend.services.embedding_service import generate_embeddings
from backend.services.qdrant_service import store_document_embeddings, delete_document
from backend.services.utils import adaptive_sentence_chunks
from tools.load_doc import load_doc_from_api
import uuid
import requests
from pathlib import Path
import logging
logger = logging.getLogger("load_doc")

# def process_document(file_path: str,doc_id:int,meeting_id:int, update: bool = False , additional_metadata:Dict[str,Any]=None) -> bool:
def process_document(meeting_id:int, doc_type:DocType ,file_path:str,doc_id:str, update: bool = False , additional_metadata:Dict[str,Any]=None) -> bool:

    """
    Process a document file with adaptive, sentence-aware chunking and rich metadata.
    Steps:
      1) Extract text and tables (page-aware for PDFs)
      2) Preprocess text
      3) Chunk by sentence boundaries (adaptive)
      4) Generate embeddings (handled downstream if not provided)
      5) Store chunks + metadata in Qdrant
    """
    try:

        filename = Path(file_path).name


        if filename.startswith('~$'):
            print(f"Skipping temporary Office file: {file_path}")
            return False


        ext = os.path.splitext(file_path)[1].lower()

        if update:
            delete_document(file_path)
        
        document_name = os.path.basename(file_path)
        all_chunks: List[str] = []
        metadatas: List[Dict[str, Any]] = []



        # Prepare base metadata to be added to other chunks added vinayak
        base_meta = {}
        if doc_id:
            base_meta["doc_id"]=doc_id
        if meeting_id:
            base_meta["meeting_id"] = meeting_id
        if additional_metadata:
            base_meta.update(additional_metadata)


        if ext == '.pdf':
            

            # Page-aware processing to enable citations with page numbers
            page_texts = extract_pdf_pages(file_path)
            logger.info(f"xxxxxx {len(page_texts)}")

            if not page_texts:
                print(f"No text extracted from {file_path}")
                return False
            for i, page_text in enumerate(page_texts, start=1):
                clean_text = preprocess_text(page_text)
                chunks = adaptive_sentence_chunks(clean_text, max_tokens=400, overlap_sentences=1)
                for idx, ch in enumerate(chunks):
                    if ch.strip():  # Skip empty chunks
                        all_chunks.append(ch)
                        chunk_meta = {
                            "chunk_id": str(uuid.uuid4()),
                            "section_title": None,
                            "page_number": i,
                            "document_hierarchy": [],
                            "source": document_name,
                            "is_table": False,
                            **base_meta #add new meta for meeting linking vinayak
                        }
                        metadatas.append(chunk_meta)

            # Extract tables from PDF
            table_chunks, table_metadatas = extract_pdf_tables(file_path, document_name)

            #add base meta to table meta
            for table_meta in table_metadatas:
                table_meta.update(base_meta)

            all_chunks.extend(table_chunks)
            metadatas.extend(table_metadatas)

      
        elif ext in ['.docx', '.doc']:
            # Extract text
            text = extract_text_from_docx(file_path)
            if text:
                clean_text = preprocess_text(text)
                chunks = adaptive_sentence_chunks(clean_text, max_tokens=400, overlap_sentences=1)
                for ch in chunks:
                    if ch.strip():
                        first_sentence = ch.split('.')[:1][0].strip() if '.' in ch else None
                        all_chunks.append(ch)
                        chunk_meta = {
                            "chunk_id": str(uuid.uuid4()),
                            "section_title": first_sentence if first_sentence else None,
                            "page_number": None,
                            "document_hierarchy": [],
                            "source": document_name,
                            "is_table": False,
                            **base_meta #add new meta for meeting linking vinayak
                        }

                        metadatas.append(chunk_meta)

            # Extract tables from DOCX
            table_chunks, table_metadatas = extract_docx_tables(file_path, document_name)

            #add base meta to table meta
            for table_meta in table_metadatas:
                table_meta.update(base_meta)

            all_chunks.extend(table_chunks)
            metadatas.extend(table_metadatas)

        elif ext in ['.xlsx', '.xls']:
            # Extract tables from Excel (sheets as tables)
            table_chunks, table_metadatas = extract_excel_tables(file_path, document_name)

            #add base meta to table meta
            for table_meta in table_metadatas:
                table_meta.update(base_meta)

            all_chunks.extend(table_chunks)
            metadatas.extend(table_metadatas)

        else:
            # Single-text processing for other formats
            text = extract_text(file_path)
            if not text:
                print(f"No text extracted from {file_path}")
                return False
            clean_text = preprocess_text(text)
            chunks = adaptive_sentence_chunks(clean_text, max_tokens=400, overlap_sentences=1)
            for ch in chunks:
                if ch.strip():
                    first_sentence = ch.split('.')[:1][0].strip() if '.' in ch else None
                    all_chunks.append(ch)

                    chunk_meta = {
                        "chunk_id": str(uuid.uuid4()),
                        "section_title": first_sentence if first_sentence else None,
                        "page_number": None,
                        "document_hierarchy": [],
                        "source": document_name,
                        "is_table": False,
                        **base_meta  # add new meta for meeting linking vinayak
                    }

                    metadatas.append(chunk_meta)

        if not all_chunks:
            print(f"No chunks generated for {file_path}")
            return False

        # Store in Qdrant (embeddings computed inside if not provided)
        store_document_embeddings(file_path, document_name, all_chunks, embeddings=None, metadatas=metadatas)

        print(f"Successfully processed document: {file_path}")
        return True

     

    except Exception as e:
        print(f"Error processing document : {str(e)}")
        return False

def remove_document(file_path: str) -> bool:
    """
    Remove a document's embeddings from Qdrant
    
    Args:
        file_path: Path to the document
        
    Returns:
        bool: Success or failure
    """
    try:
        delete_document(file_path)
        print(f"Successfully removed document: {file_path}")
        return True
    except Exception as e:
        print(f"Error removing document {file_path}: {str(e)}")
        return False

def extract_text(file_path: str) -> str:
    """
    Extract text from various file types
    
    Args:
        file_path: Path to the file
        
    Returns:
        str: Extracted text
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        # PDF files
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
            
        # Word documents
        elif file_extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
            
        # Excel files
        elif file_extension in ['.xlsx', '.xls']:
            return extract_text_from_excel(file_path)
            
        # CSV files
        elif file_extension == '.csv':
            return extract_text_from_csv(file_path)
            
        # Text files
        elif file_extension in ['.txt', '.md', '.json']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
                
        else:
            print(f"Unsupported file type: {file_extension}")
            return ""
            
    except Exception as e:
        print(f"Error extracting text from {file_path}: {str(e)}")
        return ""

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using both PyMuPDF and pdfplumber for better results"""
    text = ""
    
    # Try PyMuPDF first
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        print(f"PyMuPDF extraction failed: {str(e)}")
    
    # If PyMuPDF didn't get much text, try pdfplumber
    if len(text.strip()) < 100:
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
        except Exception as e:
            print(f"pdfplumber extraction failed: {str(e)}")
    
    return text





def extract_pdf_pages(file_path: str) -> List[str]:
    """Return list of per-page texts to enable page-number metadata."""
    pages: List[str] = []
    try:
        # Try PyMuPDF
        doc = fitz.open(file_path)
        for page in doc:
            pages.append(page.get_text() or "")
        doc.close()
    except Exception:
        pages = []
    # Fallback to pdfplumber if PyMuPDF failed or produced empty
    if not pages or sum(len(p.strip()) for p in pages) < 50:
        try:
            with pdfplumber.open(file_path) as pdf:
                pages = [(p.extract_text() or "") for p in pdf.pages]
        except Exception:
            pass
    return pages

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from Word documents"""
    doc = docx.Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def extract_text_from_excel(file_path: str) -> str:
    """Extract text from Excel files"""
    df = pd.read_excel(file_path, sheet_name=None)
    text = ""
    
    for sheet_name, sheet_df in df.items():
        text += f"Sheet: {sheet_name}\n"
        text += sheet_df.to_string(index=False) + "\n\n"
    
    return text

def extract_text_from_csv(file_path: str) -> str:
    """Extract text from CSV files"""
    df = pd.read_csv(file_path)
    return df.to_string(index=False)

def preprocess_text(text: str) -> str:
    """
    Preprocess and clean text
    
    Args:
        text: Raw text
        
    Returns:
        str: Cleaned text
    """
    # Remove HTML tags
    text = re.sub(r'<.*?>', ' ', text)
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters but keep punctuation
    text = re.sub(r'[^\w\s.,;:!?\'"-]', ' ', text)
    
    # Normalize whitespace
    text = text.strip()
    
    return text

def is_qa_document(text: str) -> bool:
    """
    Detect if document is Q&A/FAQ format
    
    Args:
        text: Document text
        
    Returns:
        bool: True if document appears to be Q&A format
    """
    qa_indicators = [
        'frequently asked questions',
        'faq',
        'q&a',
        'question:',
        'answer:',
        'q.',
        'a.',
    ]
    
    text_lower = text.lower()[:2000]  # Check first 2000 chars
    
    # Count question marks in first portion
    question_count = text[:3000].count('?')
    
    # Check for FAQ indicators
    has_qa_format = any(indicator in text_lower for indicator in qa_indicators)
    
    # If has FAQ indicators or many questions, likely Q&A format
    return has_qa_format or question_count >= 5


def split_qa_document(text: str) -> List[str]:
    """
    Split Q&A document keeping question-answer pairs together
    
    Args:
        text: Q&A document text
        
    Returns:
        List[str]: List of Q&A chunks
    """
    chunks = []
    
    # Split by common Q&A patterns
    # Pattern 1: "Question text?\nAnswer text"
    import re
    
    # Try to split by questions (sentences ending with ?)
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    current_chunk = ""
    current_question = ""
    word_count = 0
    max_chunk_words = 500  # Slightly larger for Q&A to keep pairs together
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        
        sentence_words = len(sentence.split())
        
        # If sentence is a question, start new Q&A pair
        if '?' in sentence:
            # If we have accumulated content and adding this would be too large, save chunk
            if current_chunk and word_count + sentence_words > max_chunk_words:
                chunks.append(current_chunk.strip())
                current_chunk = ""
                word_count = 0
            
            current_question = sentence
            current_chunk += "\n\n" + sentence if current_chunk else sentence
            word_count += sentence_words
        else:
            # This is likely an answer or continuation
            current_chunk += " " + sentence
            word_count += sentence_words
            
            # If chunk is getting large, split after complete Q&A pair
            if word_count >= max_chunk_words and '?' in current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""
                word_count = 0
    
    # Add remaining content
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    # Filter out very small chunks (less than 50 words)
    chunks = [c for c in chunks if len(c.split()) >= 50]
    
    return chunks if chunks else [text]  # Fallback to full text if no chunks created


def split_text_into_chunks(text: str, chunk_size: int = 400, overlap: int = 50) -> List[str]:
    """
    Intelligently split text into chunks based on document type

    Args:
        text: Text to split
        chunk_size: Target size of each chunk (in tokens, approximated by words)
        overlap: Number of tokens to overlap between chunks

    Returns:
        List[str]: List of text chunks
    """
    # Detect if this is a Q&A document
    if is_qa_document(text):
        print("Detected Q&A format document - using specialized chunking")
        return split_qa_document(text)

    # Standard chunking for non-Q&A documents
    words = text.split()

    if len(words) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk = ' '.join(words[start:end])
        chunks.append(chunk)
        start += chunk_size - overlap

    return chunks


def extract_pdf_tables(file_path: str, document_name: str) -> Tuple[List[str], List[Dict[str, Any]]]:
    """Extract tables from PDF and convert to Markdown chunks."""
    chunks = []
    metadatas = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables()
                for table_idx, table in enumerate(tables):
                    if table:
                        # Convert table to Markdown
                        markdown_table = table_to_markdown(table)
                        if markdown_table.strip():
                            chunks.append(markdown_table)
                            metadatas.append({
                                "chunk_id": str(uuid.uuid4()),
                                "section_title": f"Table {table_idx + 1}",
                                "page_number": page_num,
                                "document_hierarchy": [],
                                "source": document_name,
                                "is_table": True,
                                "sheet": None
                            })
    except Exception as e:
        print(f"Error extracting PDF tables: {str(e)}")
    return chunks, metadatas


def extract_docx_tables(file_path: str, document_name: str) -> Tuple[List[str], List[Dict[str, Any]]]:
    """Extract tables from DOCX and convert to Markdown chunks."""
    chunks = []
    metadatas = []
    try:
        doc = docx.Document(file_path)
        for table_idx, table in enumerate(doc.tables):
            # Convert table to Markdown
            markdown_table = docx_table_to_markdown(table)
            if markdown_table.strip():
                chunks.append(markdown_table)
                metadatas.append({
                    "chunk_id": str(uuid.uuid4()),
                    "section_title": f"Table {table_idx + 1}",
                    "page_number": None,
                    "document_hierarchy": [],
                    "source": document_name,
                    "is_table": True,
                    "sheet": None
                })
    except Exception as e:
        print(f"Error extracting DOCX tables: {str(e)}")
    return chunks, metadatas


def extract_excel_tables(file_path: str, document_name: str) -> Tuple[List[str], List[Dict[str, Any]]]:
    """Extract sheets from Excel as table chunks."""
    chunks = []
    metadatas = []
    try:
        xl = pd.ExcelFile(file_path)
        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            # Convert DataFrame to Markdown table
            markdown_table = df.to_markdown(index=False)
            if markdown_table.strip():
                chunks.append(markdown_table)
                metadatas.append({
                    "chunk_id": str(uuid.uuid4()),
                    "section_title": f"Sheet: {sheet_name}",
                    "page_number": None,
                    "document_hierarchy": [],
                    "source": document_name,
                    "is_table": True,
                    "sheet": sheet_name
                })
    except Exception as e:
        print(f"Error extracting Excel tables: {str(e)}")
    return chunks, metadatas


def table_to_markdown(table: List[List[str]]) -> str:
    """Convert a list of lists (table) to Markdown format."""
    if not table:
        return ""
    # Clean table data
    cleaned_table = [[cell.strip() if cell else "" for cell in row] for row in table]
    # Create header
    header = "| " + " | ".join(cleaned_table[0]) + " |"
    separator = "|" + "|".join(["---"] * len(cleaned_table[0])) + "|"
    # Create rows
    rows = ["| " + " | ".join(row) + " |" for row in cleaned_table[1:]]
    return "\n".join([header, separator] + rows)


def docx_table_to_markdown(table) -> str:
    """Convert python-docx table to Markdown format."""
    if not table.rows:
        return ""
    # Extract rows
    rows = []
    for row in table.rows:
        row_data = [cell.text.strip() for cell in row.cells]
        rows.append(row_data)
    return table_to_markdown(rows)



def fetch_documents_from_service(meeting_id: int, doc_type: str) -> List[dict]:
    """
    Fetch documents from the external document service.
    """
    try:
        url = os.getenv("SURV_PORTAL_URL", "https://survportaluat.nse.co.in/surv/document-service/document/getdocument")
        response = requests.get(
            url,
            params={"meetingId": meeting_id, "categoryId": doc_type},
            timeout=60, verify=False
        )
        logger.info(f"Following on URL {response.url}")

        if response.status_code != 200:
            logger.error(f"Document service failed with status {response.status_code}")
            return []

        json_response = response.json()
        results = json_response.get("result", [])
        
        parsed_docs = []
        for result in results:
            # Extract base64 and decode
            b64_str = result.get("base64", "")
            decoded_bytes = base64.b64decode(b64_str) if b64_str else b""
            
            # Prepare doc dict with standard keys
            doc_info = {
                "bytes": decoded_bytes,
                "filename": result.get("documentName"),
                "doc_id": result.get("id"),
            }
            # Merge all other keys from result for metadata (excluding base64 to save memory)
            for k, v in result.items():
                if k not in ["base64", "documentName", "id"]:
                    doc_info[k] = v
            
            parsed_docs.append(doc_info)

        return parsed_docs
    except Exception as e:
        logger.error(f"Error fetching documents: {e}")
        return []

def save_single_document_to_disk(doc: dict) -> Tuple[str, int]:
    import tempfile
    save_dir = Path(tempfile.gettempdir()) / "surv_documents"
    save_dir.mkdir(parents=True, exist_ok=True)

    file_path = (save_dir / doc["filename"]).resolve()

    with open(file_path, "wb") as f:
        f.write(doc["bytes"])

    logger.info(f"Document saved at {file_path}")

    return str(file_path), doc["doc_id"]
